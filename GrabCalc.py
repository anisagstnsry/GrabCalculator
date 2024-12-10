# -*- coding: utf-8 -*-
"""
Created on Sat Jan 20 12:32:43 2024

@author: 
"""

import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import base64

data_employee = pd.read_excel("master_data.xlsx", sheet_name = "Employee Data")
data_rules = pd.read_excel("master_data.xlsx", sheet_name = "Rules Data")
data_dictionary = pd.read_excel("master_data.xlsx", sheet_name = "Dictionary", index_col = "Variable")
tarifperkm = data_dictionary.loc["tarif_perkm", "Value"]
tarifminimal = data_dictionary.loc["tarif_minimal", "Value"]


def write():
    st.title("Grab Reimbursement Calculator")
    name = st.selectbox("Nama", data_employee["Full Name"])
    date_trip = st.date_input("Tanggal")
    list_waktu = ["Office hour 06.00 - 19.30", "Non-Office hour 19.30 - 06.00"]
    kategori_waktu = st.selectbox("Pilih waktu berangkat", list_waktu)
    condition_trip = st.selectbox("Pilih kondisi", ("TIDAK HUJAN DAN TIDAK FORCE MAJOURE", "HUJAN", "FORCE MAJOURE"))
    people_trip = st.number_input("Jumlah orang yang berangkat", min_value = 1)
    notes = st.text_input("Catatan untuk menjelaskan kondisi dan menyertakan nama tim jika lebih dari 1 orang")
    input_pid = st.text_input("Input PID atau nama agenda, contoh: KOM 2024 atau POC Danamon")
    list_rute = list(data_rules["Rute"].unique())
    route = st.selectbox("Pilih rute", list_rute)
    pilihan_akses = st.selectbox("Apakah akses transportasi umum terbatas?", ("Ya, e.g. hanya bisa dilalui oleh mobil lewat tol", "Tidak, e.g. masih ada akses kereta/bis"))
    jarak = st.number_input("Jarak perjalanan (KM)", min_value = 1)
    if jarak < 20:
        kategori_jarak = "Kurang dari 20 KM"
    elif jarak < 30:
        kategori_jarak = "21-30 KM"
    else: 
        kategori_jarak = "Lebih dari 30 KM"
    st.write(kategori_jarak)
    list_tipegrab = ["Grab Personal (Bike)", "Grab Personal (Car)"]
    tipe_grab = st.selectbox("Tipe transportasi", list_tipegrab)
    total_payment = st.number_input("Total yang dibayarkan (Rp)", min_value = 10000)
    
    if people_trip > 1:
        kategori_policy = "Group"
    else:
        if condition_trip == "HUJAN":
            kategori_policy = "Rainy"
        elif condition_trip == "FORCE MAJOURE":
            kategori_policy = "Force Majoure"
        elif kategori_waktu == "Non-Office hour 19.30 - 06.00":
            kategori_policy = "Non-Office hour 19.30 - 06.00"
        elif pilihan_akses == "Ya, e.g. hanya bisa dilalui oleh mobil lewat tol":
            kategori_policy = "Akses Terbatas"
        elif tipe_grab == "Grab Personal (Car)":
            kategori_policy = "Route Efficiency"
        else:
            kategori_policy = "Normal"

    data_kondisi = data_rules[data_rules["Kondisi"] == kategori_policy]
    data_kondisi = data_kondisi[data_kondisi["Kategori Jarak"] == kategori_jarak]
    data_kondisi = data_kondisi[data_kondisi["Rute"] == route]
    
    data_kondisi.reset_index(inplace = True)
    coverage_type = data_kondisi.loc[0, "Coverage"]
    if data_kondisi.loc[0, "Coverage"] == "Full":
        covered = total_payment
    elif data_kondisi.loc[0, "Coverage"] == "Not Covered":
        covered = 0
    elif data_kondisi.loc[0, "Coverage"] == "Rumus 1":
        if route == "Office - Client" or route == "Client - Office":
            jarak_adi = jarak
        else:
            jarak_adi = st.number_input("Jarak perjalanan (KM) dari kantor client ke kantor ADI atau sebaliknya", min_value = 0)
        if jarak < jarak_adi:
            covered = 0
        else:    
            covered = jarak_adi*tarifperkm + tarifminimal
    elif data_kondisi.loc[0, "Coverage"] == "Rumus 2":
        jarak_client = st.number_input("Jarak perjalanan (KM) dari/ke kantor client", min_value = 0)
        covered = jarak_client*tarifperkm + tarifminimal
        
    button_click = st.button("Calculate")
    if button_click:
        st.write("Kondisi Perjalanan: {}".format(kategori_policy))
        st.write("Nominal yang akan direimburse kantor adalah {}".format(covered))
    
        document = Document()
        document.add_heading('Grab Reimbursement', 0)
        document.add_paragraph('Nama: {}'.format(name))
        document.add_paragraph('Tanggal: {}'.format(date_trip))
        document.add_paragraph('PID/Agenda: {}'.format(input_pid))
        document.add_paragraph('Rute: {}'.format(route))
        document.add_paragraph('Catatan: {}'.format(notes))
        document.add_paragraph('Jarak: {} KM'.format(jarak))
        document.add_paragraph('Kategori Jarak: {}'.format(kategori_jarak))
        document.add_paragraph('Biaya tertagih: Rp {}'.format(f"{total_payment:,}"))
        document.add_paragraph('Moda yang digunakan: {}'.format(tipe_grab))
        document.add_paragraph('Total biaya yang dapat direimburse: Rp {}'.format(f"{covered:,}"))
        
        f = BytesIO()
        document.save(f)
        
        def download_pdf(doc):
            doc_str = base64.b64encode(doc.getvalue()).decode()
            href = f'<a href="data:file/txt;base64,{doc_str}" download="GrabCalc_{name}.docx">Download docx</a>'
            return href
    
        st.markdown(download_pdf(f), unsafe_allow_html=True)

    
if __name__ == '__main__':
    write()
